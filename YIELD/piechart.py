# import matplotlib.pyplot as plt
#
# # Crop data with approximate median values
# crops = [
#     'Wheat', 'Maize', 'Coriander', 'Fennel', 'Garlic', 'Tomato', 'Fenugreek', 'Pulses', 'Guava', 'Pomegranate',
#     'Sugarcane', 'Citrus', 'Cotton', 'Bajra', 'Oilseeds', 'Onion', 'Mustard', 'Mango', 'Chilli', 'Gram', 'Cumin',
#     'Barley', 'Opium'
# ]
# median_yields = [
#     40, 38, 36, 37, 41, 42, 35, 39, 38, 37, 36, 39, 38, 37, 36, 39, 40, 41, 39, 38, 37, 36, 35
# ]
#
# # Create a pie chart
# plt.figure(figsize=(10, 8))
# plt.pie(median_yields, labels=crops, autopct='%1.1f%%', startangle=140)
# plt.axis('equal')
# plt.title('Crop Yield Distribution')
# plt.show()



# import matplotlib.pyplot as plt
# import numpy as np
#
# # Data
# irrigation_methods = ['Sprinkler Irrigation', 'Drip Irrigation', 'Tube Well', 'Canal Irrigation']
# median_values = [37, 35, 36, 36]
# q1_values = [30, 30, 30, 30]
# q3_values = [45, 43, 44, 42]
# min_values = [20, 22, 25, 22]
# max_values = [60, 58, 55, 56]
#
# # Calculate mean values as an example
# mean_values = [(q1 + median + q3) / 3 for q1, median, q3 in zip(q1_values, median_values, q3_values)]
#
# # Bar Graph for Mean Yield
# plt.figure(figsize=(10, 6))
# plt.bar(irrigation_methods, mean_values, color=['blue', 'lightblue', 'orange', 'tan'])
# plt.xlabel('Irrigation Method')
# plt.ylabel('Mean Yield (Quintals)')
# plt.title('Mean Yield by Irrigation Method')
# plt.show()
from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('Project Block Diagram', level=1)

# Define the blocks and their content
blocks = {
    "Data Collection and Preprocessing": [
        "Gathering all software code and documents from the company.",
        "Cleaning the data to ensure it is free of errors and inconsistencies."
    ],
    "Natural Language Processing (NLP)": [
        "Utilizing NLP algorithms to understand and interpret the code.",
        "Generating human-readable documentation from the code using NLP models."
    ],
    "Interactive Query System": [
        "Implementing a chatbot using frameworks like Rasa or HuggingFace API.",
        "Enabling the chatbot to understand user queries and provide relevant code segments and documentation."
    ],
    "Issue Detection and Suggestions": [
        "Developing and integrating algorithms that can detect issues in the code.",
        "Incorporating these capabilities into the chatbot for real-time problem identification and suggestion."
    ],
    "Testing and Validation": [
        "Testing the system with real data to validate its accuracy and usability.",
        "Ensuring the generated documentation and interactive query responses are correct and helpful."
    ],
    "Expected Outcomes": [
        "Developing an automated system for generating coding documentation.",
        "Creating an interactive query resolver.",
        "Establishing a fast and efficient method for document creation and query resolution."
    ],
    "Conclusion": [
        "Emphasizing the improvement in efficiency and accuracy of code documentation and retrieval.",
        "Highlighting how the project will enhance software maintenance and debugging."
    ]
}

# Add blocks to the document
for block_title, content in blocks.items():
    doc.add_heading(block_title, level=2)
    for item in content:
        p = doc.add_paragraph(f'- {item}')
        p.style.font.size = Pt(12)

# Save the document to a valid local path
file_path = "C:/Users/Ayush/Documents/Project_Block_Diagram.docx"
doc.save(file_path)

print(f"Document saved to {file_path}")

